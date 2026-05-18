from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from app.formats.text import _page_text
from app.ocr.base import OcrResult


def write_docx(result: OcrResult, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / "output.docx"
    doc = Document()
    for i, page in enumerate(result.pages):
        if i > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(f"Page {page.page_number}", level=2)
        for line in _page_text(page).splitlines():
            if line.strip():
                doc.add_paragraph(line)
    doc.save(str(out))
    return out
